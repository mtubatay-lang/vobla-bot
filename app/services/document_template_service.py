"""Сервис генерации документов из DOCX-шаблонов: загрузка конфига, заполнение плейсхолдеров, извлечение полей через OpenAI."""

import io
import json
import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from openai import OpenAI

from app.config import TEMPLATES_DIR, OPENAI_API_KEY, OPENAI_MODEL, OPENAI_TIMEOUT
from app.services.document_processor import extract_text

logger = logging.getLogger(__name__)

# Placeholder format in DOCX: {{field_key}}
PLACEHOLDER_PATTERN = re.compile(r"\{\{(\w+)\}\}")

# --- Config loading ---


def load_templates_config() -> Dict[str, Any]:
    """Загружает конфигурацию шаблонов из templates/config.json."""
    config_path = TEMPLATES_DIR / "config.json"
    if not config_path.is_file():
        logger.warning("[DOC_TEMPLATE] config.json не найден: %s", config_path)
        return {"templates": []}
    try:
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError) as e:
        logger.exception("[DOC_TEMPLATE] Ошибка чтения config.json: %s", e)
        return {"templates": []}


def get_templates_list() -> List[Dict[str, Any]]:
    """Возвращает список шаблонов из конфига."""
    config = load_templates_config()
    return config.get("templates", []) or []


def get_template_by_id(template_id: str) -> Optional[Dict[str, Any]]:
    """Возвращает шаблон по id или None."""
    for t in get_templates_list():
        if t.get("id") == template_id:
            return t
    return None


def get_template_path(template: Dict[str, Any]) -> Path:
    """Возвращает полный путь к файлу шаблона."""
    filename = template.get("filename", "")
    return TEMPLATES_DIR / filename


# --- DOCX filling (paragraphs, tables, headers/footers; handle split runs) ---


def _replace_in_paragraph(paragraph, replacements: Dict[str, str]) -> None:
    """Заменяет плейсхолдеры в параграфе. Объединяет текст runs для поиска {{key}}."""
    full_text = paragraph.text
    if not full_text or "{{" not in full_text:
        return
    for key, value in replacements.items():
        placeholder = "{{" + key + "}}"
        if placeholder in full_text:
            # Replace in each run that contains part of the placeholder
            # Strategy: rebuild runs by doing full-text replace then assign to first run, clear others
            new_text = full_text.replace(placeholder, str(value))
            if new_text != full_text:
                # Clear all runs and set full text in first run to preserve formatting
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ""


def _replace_in_cell(cell, replacements: Dict[str, str]) -> None:
    """Заменяет плейсхолдеры во всех параграфах ячейки таблицы."""
    for paragraph in cell.paragraphs:
        _replace_in_paragraph(paragraph, replacements)


def _replace_in_document(doc: Document, replacements: Dict[str, str]) -> None:
    """Проходит по всем параграфам, таблицам и колонтитулам, заменяет {{key}}."""
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_cell(cell, replacements)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)


def _replace_literal_in_paragraph(paragraph, literal: str, placeholder: str) -> None:
    """Заменяет literal на placeholder в параграфе (для подготовки шаблона)."""
    full_text = paragraph.text
    if not full_text or literal not in full_text:
        return
    new_text = full_text.replace(literal, placeholder)
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = new_text
        else:
            run.text = ""


def replace_literals_with_placeholders_in_document(
    doc: Document,
    replacements: List[tuple],
) -> None:
    """
    Заменяет литералы на плейсхолдеры в документе (для создания шаблона из примера).
    replacements: [(literal_str, placeholder_str), ...]
    """
    for para in doc.paragraphs:
        for literal, placeholder in replacements:
            _replace_literal_in_paragraph(para, literal, placeholder)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for literal, placeholder in replacements:
                        _replace_literal_in_paragraph(para, literal, placeholder)
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    for literal, placeholder in replacements:
                        _replace_literal_in_paragraph(para, literal, placeholder)


def fill_docx_template(template_path: Path, data: Dict[str, str]) -> bytes:
    """
    Заполняет DOCX-шаблон значениями из data.
    data: { "field_key": "value", ... }
    Возвращает байты готового DOCX.
    """
    doc = Document(template_path)
    # Ensure all values are strings
    replacements = {k: (v if v is not None else "") for k, v in data.items()}
    _replace_in_document(doc, replacements)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# --- Extract fields from user input (text / image / PDF / DOCX) via OpenAI ---

_openai_client: Optional[OpenAI] = None


def _get_openai_client() -> OpenAI:
    global _openai_client
    if _openai_client is None:
        if not OPENAI_API_KEY:
            raise RuntimeError("OPENAI_API_KEY не задан")
        _openai_client = OpenAI(api_key=OPENAI_API_KEY, timeout=OPENAI_TIMEOUT)
    return _openai_client


def _build_extraction_prompt(field_specs: List[Dict[str, Any]]) -> str:
    """Формирует описание полей для промпта извлечения."""
    lines = []
    for f in field_specs:
        key = f.get("key", "")
        label = f.get("label", key)
        example = f.get("example", "")
        line = f"- {key}: {label}"
        if example:
            line += f" (пример: {example})"
        lines.append(line)
    return "\n".join(lines)


def extract_fields_from_text(
    text: str,
    field_specs: List[Dict[str, Any]],
) -> Dict[str, str]:
    """
    Извлекает значения полей из текста с помощью OpenAI.
    field_specs: список из config (key, label, example).
    Возвращает { "field_key": "value", ... }. Пустые или не найденные — пустая строка.
    """
    if not text or not text.strip():
        return {f.get("key", ""): "" for f in field_specs}

    fields_desc = _build_extraction_prompt(field_specs)
    keys = [f.get("key", "") for f in field_specs]
    keys_str = ", ".join(f'"{k}"' for k in keys)

    system_prompt = (
        "Ты извлекаешь структурированные данные из текста пользователя (договоры, реквизиты, даты). "
        "Отвечай ТОЛЬКО валидным JSON без markdown и комментариев. "
        f"Ключи объекта должны быть ровно: {keys_str}. "
        "Если какое-то поле не найдено в тексте — используй пустую строку \"\" для этого ключа. "
        "Даты пиши в формате «день месяц год» (например: 22 февраля 2026). "
        "Реквизиты (ИНН, адрес, р/с, БИК и т.д.) объединяй в одно поле customer_details с переносами строк."
    )
    user_prompt = (
        f"Поля для извлечения:\n{fields_desc}\n\n"
        f"Текст пользователя:\n{text[:12000]}\n\n"
        "Верни JSON-объект с ключами из списка полей."
    )

    try:
        client = _get_openai_client()
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        raw = (resp.choices[0].message.content or "").strip()
        if not raw:
            return {k: "" for k in keys}
        data = json.loads(raw)
        result = {}
        for k in keys:
            result[k] = str(data.get(k, "")).strip() if data.get(k) else ""
        return result
    except (json.JSONDecodeError, KeyError) as e:
        logger.warning("[DOC_TEMPLATE] Ошибка разбора JSON от OpenAI: %s", e)
        return {k: "" for k in keys}
    except Exception as e:
        logger.exception("[DOC_TEMPLATE] Ошибка извлечения полей: %s", e)
        return {k: "" for k in keys}


def extract_fields_from_image(
    image_bytes: bytes,
    field_specs: List[Dict[str, Any]],
    mime_type: str = "image/jpeg",
) -> Dict[str, str]:
    """
    Извлекает значения полей из изображения (скриншот/фото) через OpenAI Vision.
    """
    import base64

    keys = [f.get("key", "") for f in field_specs]
    fields_desc = _build_extraction_prompt(field_specs)
    keys_str = ", ".join(f'"{k}"' for k in keys)

    system_prompt = (
        "Ты извлекаешь структурированные данные из изображения (скриншот, фото документа, реквизиты). "
        "Распознай весь текст на изображении и извлеки указанные поля. "
        "Отвечай ТОЛЬКО валидным JSON без markdown и комментариев. "
        f"Ключи объекта: {keys_str}. "
        "Если поле не найдено — используй пустую строку \"\"."
    )
    user_content: List[Any] = [
        {
            "type": "text",
            "text": f"Поля для извлечения:\n{fields_desc}\n\nВерни JSON-объект с этими ключами по содержимому изображения.",
        },
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:{mime_type};base64,{base64.b64encode(image_bytes).decode('ascii')}"
            },
        },
    ]

    try:
        client = _get_openai_client()
        resp = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.1,
            response_format={"type": "json_object"},
        )
        raw = (resp.choices[0].message.content or "").strip()
        if not raw:
            return {k: "" for k in keys}
        data = json.loads(raw)
        result = {}
        for k in keys:
            result[k] = str(data.get(k, "")).strip() if data.get(k) else ""
        return result
    except (json.JSONDecodeError, KeyError) as e:
        logger.warning("[DOC_TEMPLATE] Ошибка разбора JSON (Vision): %s", e)
        return {k: "" for k in keys}
    except Exception as e:
        logger.exception("[DOC_TEMPLATE] Ошибка извлечения полей из изображения: %s", e)
        return {k: "" for k in keys}


def extract_fields_from_file(
    file_bytes: bytes,
    filename: str,
    field_specs: List[Dict[str, Any]],
) -> Dict[str, str]:
    """
    Универсальный вход: по расширению извлекаем текст (PDF/DOCX/TXT) и передаём в extract_fields_from_text.
    """
    try:
        text = extract_text(file_bytes, filename)
        return extract_fields_from_text(text, field_specs)
    except Exception as e:
        logger.exception("[DOC_TEMPLATE] Ошибка извлечения текста из файла %s: %s", filename, e)
        return {f.get("key", ""): "" for f in field_specs}


# --- Template upload: extract placeholders from DOCX, save file, update config ---


def extract_placeholder_keys_from_docx(file_bytes: bytes) -> List[str]:
    """
    Извлекает уникальные ключи плейсхолдеров {{key}} из DOCX (параграфы и таблицы).
    """
    doc = Document(io.BytesIO(file_bytes))
    keys_seen: set = set()
    for paragraph in doc.paragraphs:
        for m in PLACEHOLDER_PATTERN.finditer(paragraph.text):
            keys_seen.add(m.group(1))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for m in PLACEHOLDER_PATTERN.finditer(paragraph.text):
                        keys_seen.add(m.group(1))
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for paragraph in header_footer.paragraphs:
                    for m in PLACEHOLDER_PATTERN.finditer(paragraph.text):
                        keys_seen.add(m.group(1))
    return sorted(keys_seen)


def _slug(s: str) -> str:
    """Делает slug из названия для id и имени файла."""
    s = (s or "").strip().lower()
    result = []
    for c in s:
        if c.isalnum() or c in " -_":
            result.append(c if c != " " else "_")
    return "".join(result).strip("_") or "template"


def save_template_file(filename: str, file_bytes: bytes) -> Path:
    """Сохраняет файл шаблона в TEMPLATES_DIR. Возвращает путь к файлу."""
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    path = TEMPLATES_DIR / filename
    path.write_bytes(file_bytes)
    return path


def add_template_to_config(
    template_id: str,
    name: str,
    filename: str,
    field_keys: List[str],
) -> None:
    """
    Добавляет шаблон в config.json. template_id и filename должны быть заданы.
    field_keys — список ключей полей (key и label будут одинаковые или label из key).
    """
    config = load_templates_config()
    templates = list(config.get("templates", []) or [])
    fields = [{"key": k, "label": k.replace("_", " ").title()} for k in field_keys]
    templates.append({
        "id": template_id,
        "name": name,
        "filename": filename,
        "fields": fields,
    })
    config["templates"] = templates
    config_path = TEMPLATES_DIR / "config.json"
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    with open(config_path, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)
