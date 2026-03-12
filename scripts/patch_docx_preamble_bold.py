#!/usr/bin/env python3
"""
Разбивает параграф с {{customer_preamble}} на два run: только подставляемое значение жирным,
«, именуемый в дальнейшем "Заказчик"» и далее — обычным шрифтом.
Запуск из корня проекта: python scripts/patch_docx_preamble_bold.py
"""

import sys
from pathlib import Path

project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from docx import Document

from app.config import TEMPLATES_DIR

PLACEHOLDER = "{{customer_preamble}}"


def _split_paragraph_runs(paragraph):
    """Если параграф содержит PLACEHOLDER и запятую после него — разбить на два run: плейсхолдер (жирный), остальное (не жирный)."""
    text = paragraph.text
    if PLACEHOLDER not in text or text.index(PLACEHOLDER) + len(PLACEHOLDER) >= len(text):
        return False
    after_placeholder = text[text.index(PLACEHOLDER) + len(PLACEHOLDER) :]
    if not after_placeholder.startswith(","):
        return False
    # Один run с плейсхолдером (жирный), второй с остальным текстом (не жирный)
    runs = list(paragraph.runs)
    if not runs:
        return False
    runs[0].text = PLACEHOLDER
    runs[0].bold = True
    if len(runs) > 1:
        runs[1].text = after_placeholder
        runs[1].bold = False
        for r in runs[2:]:
            r.text = ""
    else:
        paragraph.add_run(after_placeholder).bold = False
    return True


def _process_doc(doc):
    changed = 0
    for p in doc.paragraphs:
        if _split_paragraph_runs(p):
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if _split_paragraph_runs(p):
                        changed += 1
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for p in hf.paragraphs:
                    if _split_paragraph_runs(p):
                        changed += 1
    return changed


def main():
    for filename in ("accounting_contract_template.docx", "marketing_contract_template.docx"):
        path = TEMPLATES_DIR / filename
        if not path.is_file():
            print(f"Пропуск (файл не найден): {filename}")
            continue
        doc = Document(path)
        n = _process_doc(doc)
        doc.save(path)
        print(f"Обновлён: {filename} (параграфов разбито: {n})")


if __name__ == "__main__":
    main()
