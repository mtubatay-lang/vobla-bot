#!/usr/bin/env python3
"""
Создаёт DOCX-шаблоны с плейсхолдерами из исходных документов.
Запуск из корня проекта:
  python scripts/create_document_templates.py [путь_к_бух_дог] [путь_к_маркетинг_дог]

Если пути не указаны, копирует из templates/source/ если есть, иначе создаёт минимальные шаблоны.
"""

import sys
from pathlib import Path

# Добавляем корень проекта в path
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from docx import Document

from app.config import TEMPLATES_DIR
from app.services.document_template_service import (
    replace_literals_with_placeholders_in_document,
)


def main():
    templates_dir = TEMPLATES_DIR
    templates_dir.mkdir(parents=True, exist_ok=True)

    # Бухгалтерский договор
    accounting_source = project_root / "scripts" / "source" / "Бух дог ШАБЛОН — копия.docx"
    if len(sys.argv) > 1:
        accounting_source = Path(sys.argv[1])
    accounting_out = templates_dir / "accounting_contract_template.docx"

    if accounting_source.is_file():
        doc = Document(accounting_source)
        replace_literals_with_placeholders_in_document(
            doc,
            [
                ("015-2508", "{{contract_number}}"),
                ('"22" февраля 2026', "{{contract_date}}"),
                ("22 февраля 2026", "{{contract_date}}"),
                ('"22" августа 2025', "{{contract_date}}"),
                ("22 августа 2025", "{{contract_date}}"),
                ("«31» декабря 2026", "{{contract_end_date}}"),
                ("31 декабря 2026", "{{contract_end_date}}"),
                (
                    "Индивидуальный предприниматель, именуемый в дальнейшем \"Заказчик\"",
                    "Индивидуальный предприниматель {{customer_details}}, именуемый в дальнейшем \"Заказчик\"",
                ),
            ],
        )
        doc.save(accounting_out)
        print(f"Создан: {accounting_out}")
    else:
        _create_minimal_template(
            accounting_out,
            "Договор N {{contract_number}}\nвозмездного оказания услуг\n\n"
            "г. Казань {{contract_date}}\n\n"
            "Индивидуальный предприниматель {{customer_details}}, именуемый в дальнейшем \"Заказчик\"...\n\n"
            "Срок действия до {{contract_end_date}} года.",
        )
        print(f"Минимальный шаблон создан: {accounting_out}")

    # Маркетинговый договор
    marketing_source = project_root / "scripts" / "source" / "Новый_шаблон_договора_оказания_маркетинг_услуг.docx"
    if len(sys.argv) > 2:
        marketing_source = Path(sys.argv[2])
    marketing_out = templates_dir / "marketing_contract_template.docx"

    if marketing_source.is_file():
        doc = Document(marketing_source)
        replace_literals_with_placeholders_in_document(
            doc,
            [
                ("ДК 026/2602", "{{contract_number}}"),
                ("01 марта 2026", "{{contract_date}}"),
                ("«01» марта 2026", "{{contract_date}}"),
                (
                    "Индивидуальный предприниматель, действующий на основании",
                    "Индивидуальный предприниматель {{customer_details}}, действующий на основании",
                ),
            ],
        )
        # Таблица с датой — обрабатывается через replace выше по ячейкам
        doc.save(marketing_out)
        print(f"Создан: {marketing_out}")
    else:
        _create_minimal_template(
            marketing_out,
            "Договор возмездного оказания услуг № {{contract_number}}\n\n"
            "Индивидуальный предприниматель {{customer_details}}, действующий на основании...\n\n"
            "Дата: {{contract_date}}\nСрок: {{contract_end_date}}",
        )
        print(f"Минимальный шаблон создан: {marketing_out}")


def _create_minimal_template(path: Path, content: str) -> None:
    """Создаёт минимальный DOCX с одним параграфом и плейсхолдерами."""
    doc = Document()
    for part in content.split("\n\n"):
        doc.add_paragraph(part)
    doc.save(path)


if __name__ == "__main__":
    main()
