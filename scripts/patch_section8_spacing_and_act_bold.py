#!/usr/bin/env python3
"""
Патч шаблона: 1) пустая строка между Заказчик и Исполнитель в разделе 8;
2) заголовок «Исполнитель» жирным в акте (два run).
Запуск из корня проекта: python scripts/patch_section8_spacing_and_act_bold.py
"""

import sys
from pathlib import Path

project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root))

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from docx import Document
from app.config import TEMPLATES_DIR


def insert_empty_paragraph_after(paragraph):
    """Вставляет пустой параграф сразу после заданного."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def main():
    doc_path = TEMPLATES_DIR / "accounting_contract_template.docx"
    if not doc_path.is_file():
        print(f"Файл не найден: {doc_path}")
        return

    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    # 1. Пустая строка между Заказчик и Исполнитель в разделе 8
    for i in range(len(paragraphs) - 1):
        if paragraphs[i].text.strip() == "{{customer_requisites}}":
            next_text = paragraphs[i + 1].text.strip()
            if next_text.startswith("Исполнитель"):
                insert_empty_paragraph_after(paragraphs[i])
                print("Раздел 8: добавлена пустая строка между Заказчик и Исполнитель")
                break

    # 2. В акте параграф «Исполнитель:\n{{executor_requisites}}» — два run: «Исполнитель:\n» (жирный), «{{executor_requisites}}» (не жирный)
    for i in range(len(paragraphs)):
        if i < 170:
            continue
        p = paragraphs[i]
        if "Исполнитель:" in p.text and "{{executor_requisites}}" in p.text:
            if len(p.runs) == 1:
                p.runs[0].text = "Исполнитель:\n"
                p.runs[0].bold = True
                p.add_run("{{executor_requisites}}").bold = False
                print("Акт: заголовок «Исполнитель» выделен жирным")
            break

    doc.save(doc_path)
    print("Готово: accounting_contract_template.docx")


if __name__ == "__main__":
    main()
